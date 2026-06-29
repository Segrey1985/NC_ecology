"""Сборка титульного листа и объединённого документа «Сводный отчёт ООС».

Используется в пайплайне generate_all_chapters: после генерации глав 0/1/2/6
формирует:
  1) title.docx        — заполненный титульный лист;
  2) Сводный_отчёт_ООС.docx — Титул + Аннотация(ch0) + Глава1 + Глава2 + Глава6
     с разрывами страниц между разделами.

Зависит от docxcompose (склейка с сохранением форматирования) и
docx_template_engine.fill_docx_template (заполнение титула).
"""
from __future__ import annotations

from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_BREAK

try:
    from docxcompose.composer import Composer
    _HAS_COMPOSER = True
except Exception:  # pragma: no cover
    _HAS_COMPOSER = False

from src.templates.docx_template_engine import fill_docx_template
from src.utils.logger import logger


COMBINED_NAME = "Сводный_отчёт_ООС.docx"
TITLE_NAME = "title.docx"

# Шаблон титульного листа лежит рядом с этим модулем (src/templates/).
TITLE_TEMPLATE_PATH = Path(__file__).resolve().parent / "title_template.docx"

# Дефолты для полей титула, отсутствующих в данных проекта.
TITLE_DEFAULTS = {
    "СРО_НОМЕР": "СРО-П-082-14122009 от 22 октября 2019 г.",
    "ШИФР_ПРОЕКТА": "ОК.__.__/СТ–ООС",
    "ПОДПИСЬ_ДИРЕКТОР": "",
    "ПОДПИСЬ_ГИП": "",
    "ПОДПИСЬ_РАЗРАБОТЧИК": "",
    "ПОДПИСЬ_ПРОВЕРИЛ": "",
}


def build_title(
    template_path: Path,
    base_placeholders: dict,
    output_path: Path,
    *,
    clean_final: bool = True,
) -> Path:
    """Заполняет шаблон титульного листа данными проекта."""
    data = dict(TITLE_DEFAULTS)
    # Берём поля из base_placeholders (chapter0), если есть
    for key in (
        "РАЗРАБОТЧИК_РАЗДЕЛА", "НАИМЕНОВАНИЕ_ПРОЕКТА", "АДРЕС_ОБЪЕКТА",
        "МОЩНОСТЬ_ОБЪЕКТА_МВт", "ШИФР_ПРОЕКТА", "СРО_НОМЕР",
    ):
        val = base_placeholders.get(key)
        if val:
            data[key] = val
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fill_docx_template(
        template_path=template_path,
        data=data,
        output_docx_path=output_path,
        clean_final=clean_final,
        highlight=not clean_final,
    )
    logger.info(f"[title] Титульный лист сформирован: {output_path}")
    return output_path


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def combine_documents(parts: list[Path], output_path: Path) -> Path:
    """Склеивает список DOCX в один файл с разрывами страниц между ними."""
    parts = [p for p in parts if p and Path(p).exists()]
    if not parts:
        raise ValueError("Нет ни одного документа для объединения")

    if _HAS_COMPOSER:
        master = Document(str(parts[0]))
        composer = Composer(master)
        for p in parts[1:]:
            # разрыв страницы перед следующей частью
            _add_page_break(master)
            composer.append(Document(str(p)))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        composer.save(str(output_path))
    else:
        # Фолбэк без docxcompose: переносим элементы тела вручную
        master = Document(str(parts[0]))
        for p in parts[1:]:
            _add_page_break(master)
            sub = Document(str(p))
            for element in sub.element.body:
                master.element.body.append(deepcopy(element))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        master.save(str(output_path))

    logger.info(f"[combine] Объединённый документ сформирован: {output_path}")
    return output_path


def build_combined_report(
    *,
    title_template_path: Path,
    base_placeholders: dict,
    chapter0_docx: Path | None,
    chapter1_docx: Path | None,
    chapter2_docx: Path | None,
    chapter6_docx: Path | None,
    output_dir: Path,
) -> tuple[Path, Path]:
    """Создаёт титул и объединённый отчёт. Возвращает (title_path, combined_path).

    Главы, которых нет (None / не существуют), пропускаются.
    """
    title_path = build_title(
        template_path=title_template_path,
        base_placeholders=base_placeholders,
        output_path=output_dir / TITLE_NAME,
        clean_final=True,
    )
    parts = [title_path, chapter0_docx, chapter1_docx, chapter2_docx, chapter6_docx]
    parts = [p for p in parts if p is not None and Path(p).exists()]
    combined_path = combine_documents(parts, output_dir / COMBINED_NAME)
    return title_path, combined_path
